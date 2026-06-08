"""
Generates PanelIQ_IT_Requirements_June2026.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r'C:\Users\KarthikShanmugam\ClaudePOC\paneliq'

D_DARK_BLUE = RGBColor(0x1a, 0x3a, 0x5c)
D_TEAL      = RGBColor(0x00, 0x99, 0xcc)
D_WHITE     = RGBColor(0xff, 0xff, 0xff)
D_GRAY      = RGBColor(0x2d, 0x37, 0x48)
D_GREEN     = RGBColor(0x00, 0xa8, 0x6b)
D_LIGHT_ROW = RGBColor(0xf0, 0xf7, 0xfb)


def cell_bg(cell, hex_str):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_str)
    pr.append(shd)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = D_DARK_BLUE
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = D_TEAL
    return p


def bullet(doc, bold_part, rest=''):
    p = doc.add_paragraph(style='List Bullet')
    if bold_part:
        r = p.add_run(bold_part + (': ' if rest else ''))
        r.bold = True
        r.font.color.rgb = D_DARK_BLUE
        r.font.size = Pt(10.5)
    if rest:
        r2 = p.add_run(rest)
        r2.font.size = Pt(10.5)


def doc_table(doc, headers, rows, highlight_last=False, first_col_dark=True):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_bg(c, '1a3a5c')
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = D_WHITE
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        is_total = highlight_last and ri == len(rows) - 1
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            if is_total:
                cell_bg(c, '1a3a5c')
                r = c.paragraphs[0].add_run(val)
                r.bold = True; r.font.color.rgb = D_WHITE; r.font.size = Pt(10)
            elif first_col_dark and ci == 0 and len(headers) >= 3:
                cell_bg(c, '1a3a5c')
                r = c.paragraphs[0].add_run(val)
                r.bold = True; r.font.color.rgb = D_WHITE; r.font.size = Pt(10)
            else:
                if ri % 2 == 0:
                    cell_bg(c, 'f0f7fb')
                r = c.paragraphs[0].add_run(val)
                r.font.size = Pt(10)
    return t


def checklist_item(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run('☐  ' + text)
    r.font.size = Pt(10.5)


def code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9.5)
        r.font.color.rgb = D_GRAY


def create_word(path):
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin   = Inches(1.2)
        s.right_margin  = Inches(1.2)

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run('PanelIQ')
    r.bold = True; r.font.size = Pt(38); r.font.color.rgb = D_DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('IT Infrastructure Requirements')
    r.font.size = Pt(20); r.font.color.rgb = D_TEAL

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('For review and action by the IT / Infrastructure team')
    r.font.size = Pt(13); r.font.color.rgb = D_GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Borderless Access  |  June 2026')
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x71, 0x8a, 0xa4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Requestor: BA Panel Analytics  |  BA.PanelAnalytics@borderlessaccess.com')
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x71, 0x8a, 0xa4)

    doc.add_page_break()

    # ── 1. OVERVIEW ─────────────────────────────────────────────────────────
    h1(doc, '1.  Overview')
    doc.add_paragraph(
        'PanelIQ is an internal web application that enables BA team members to query panel survey '
        'data using plain English questions. The tool returns instant answers with charts, data tables, '
        'and AI-generated analyst summaries — without requiring SQL knowledge or relying on the '
        'Technology, Finance, or Analytics teams.'
    )
    doc.add_paragraph(
        'The application consists of a Python/FastAPI backend, a React frontend, and connects to two '
        'SQL Server databases. A Proof of Concept has been developed and validated. This document '
        'outlines everything IT needs to provision and configure for a production deployment at '
        'paneliq.borderlessaccess.com.'
    )

    # ── 2. SERVER REQUIREMENTS ───────────────────────────────────────────────
    doc.add_paragraph()
    h1(doc, '2.  Server Requirements')
    doc_table(doc,
        ['Item', 'Requirement'],
        [
            ['Operating System', 'Windows Server 2019/2022  or  Ubuntu 22.04 LTS'],
            ['CPU',              '4 vCPUs minimum'],
            ['RAM',              '8 GB minimum  (16 GB recommended)'],
            ['Disk',             '50 GB'],
            ['Network',          'Internal network access to SQL Server at 10.20.30.12'],
        ],
        first_col_dark=True
    )

    # ── 3. SOFTWARE TO INSTALL ───────────────────────────────────────────────
    doc.add_paragraph()
    h1(doc, '3.  Software to Install on Server')
    doc_table(doc,
        ['Software', 'Version', 'Purpose'],
        [
            ['Python',                   '3.11 or higher',   'Backend runtime — required at all times'],
            ['pip',                      'Latest',           'Python package manager'],
            ['Node.js',                  '18 or higher',     'Frontend build only — not required at runtime after build'],
            ['ODBC Driver for SQL Server','Version 17',      'Database connectivity from Python to SQL Server'],
            ['Nginx',                    'Latest stable',    'Reverse proxy + serving React static files'],
            ['NSSM',                     'Latest',           'Windows only — run FastAPI as a Windows Service'],
        ],
        first_col_dark=True
    )

    # ── 4. NETWORK & DNS ────────────────────────────────────────────────────
    doc.add_paragraph()
    h1(doc, '4.  Network & DNS Requirements')
    doc_table(doc,
        ['Item', 'Detail'],
        [
            ['Subdomain',         'paneliq.borderlessaccess.com'],
            ['SSL Certificate',   'Required — internal CA certificate or Let\'s Encrypt (if externally reachable)'],
            ['Port 80',           'HTTP — redirect all traffic to HTTPS'],
            ['Port 443',          'HTTPS — served by Nginx'],
            ['Port 8000',         'Internal only — FastAPI backend. Must NOT be exposed externally.'],
            ['Firewall Outbound', '10.20.30.12:1433 — SQL Server (KpiReports, existing analytics database)'],
            ['Firewall Outbound', 'api.anthropic.com:443 — Claude AI API (outbound HTTPS, required for AI functionality)'],
        ],
        first_col_dark=True
    )

    p = doc.add_paragraph()
    r = p.add_run('Note: ')
    r.bold = True; r.font.color.rgb = D_DARK_BLUE
    r2 = p.add_run(
        'The outbound connection to api.anthropic.com:443 is essential. '
        'If this is blocked by a firewall or proxy, the AI functionality will not work. '
        'Please confirm with the network team that this outbound HTTPS call is permitted.'
    )
    r2.font.size = Pt(10.5); r2.font.color.rgb = D_GRAY

    # ── 5. DATABASE REQUIREMENTS ─────────────────────────────────────────────
    doc.add_paragraph()
    h1(doc, '5.  Database Requirements')

    h2(doc, '5.1  DB 1 — KpiReports  (Existing, Read-Only)')
    doc_table(doc,
        ['Item', 'Detail'],
        [
            ['Server',      '10.20.30.12,1433'],
            ['Database',    'KpiReports'],
            ['Access Type', 'Read-only'],
            ['Views Used',  'KPISurveyData, KPIReportProjectData, KPIsurveydataExternal'],
            ['Credentials', 'Existing KPIUser credentials — please confirm these remain valid from the new server'],
        ],
        first_col_dark=True
    )

    doc.add_paragraph()
    h2(doc, '5.2  DB 2 — PanelIQ User Database  (New — IT to Create)')
    doc_table(doc,
        ['Item', 'Detail'],
        [
            ['Server',          'TBD — can be the same SQL Server instance as KpiReports'],
            ['Suggested DB Name','PanelIQ_Users  (or as per IT naming convention)'],
            ['Access Type',     'Read and write'],
            ['Tables Required', 'Schema will be provided by the BA Analytics team (see below)'],
        ],
        first_col_dark=True
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Tables needed in PanelIQ_Users:')
    r.bold = True; r.font.color.rgb = D_DARK_BLUE; r.font.size = Pt(11)

    bullet(doc, 'Users', 'Login accounts, roles, business unit assignment, password hashes')
    bullet(doc, 'QueryHistory', 'Log of questions asked by each user — used for personalised suggestions')
    bullet(doc, 'AccessRules', 'Row-level and field-level access control configuration per user/role')

    p = doc.add_paragraph()
    r = p.add_run('The BA Analytics team will provide the full SQL schema (CREATE TABLE scripts) separately.')
    r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = D_GRAY

    # ── 6. ENVIRONMENT VARIABLES ─────────────────────────────────────────────
    doc.add_paragraph()
    h1(doc, '6.  Environment Variables / Secrets')
    doc.add_paragraph(
        'The following values must be stored securely on the server as environment variables '
        'or in a secured .env file. They must NOT be stored in the application code or version control.'
    )
    code_block(doc, [
        'ANTHROPIC_API_KEY=<provided by BA Analytics>',
        '',
        '# KpiReports — read-only analytics database',
        'DB_SERVER=10.20.30.12,1433',
        'DB_NAME=KpiReports',
        'DB_USER=<read-only SQL user>',
        'DB_PASSWORD=<provided>',
        '',
        '# PanelIQ_Users — application database',
        'PANELIQ_DB_SERVER=<server address>',
        'PANELIQ_DB_NAME=PanelIQ_Users',
        'PANELIQ_DB_USER=<read-write SQL user>',
        'PANELIQ_DB_PASSWORD=<provided by IT>',
    ])

    # ── 7. NGINX CONFIGURATION ───────────────────────────────────────────────
    doc.add_paragraph()
    h1(doc, '7.  Suggested Nginx Configuration')
    doc.add_paragraph(
        'The following Nginx configuration routes all traffic through a single subdomain. '
        'Requests to /api/ are proxied to the FastAPI backend on port 8000. '
        'All other requests serve the React frontend static files.'
    )
    code_block(doc, [
        'server {',
        '    listen 80;',
        '    server_name paneliq.borderlessaccess.com;',
        '    return 301 https://$host$request_uri;',
        '}',
        '',
        'server {',
        '    listen 443 ssl;',
        '    server_name paneliq.borderlessaccess.com;',
        '',
        '    ssl_certificate     /path/to/certificate.crt;',
        '    ssl_certificate_key /path/to/certificate.key;',
        '',
        '    # Serve React frontend',
        '    root /var/www/paneliq/dist;',
        '    index index.html;',
        '',
        '    # React SPA routing',
        '    location / {',
        '        try_files $uri $uri/ /index.html;',
        '    }',
        '',
        '    # Proxy API calls to FastAPI backend',
        '    location /api/ {',
        '        proxy_pass http://127.0.0.1:8000/;',
        '        proxy_set_header Host $host;',
        '        proxy_set_header X-Real-IP $remote_addr;',
        '        proxy_read_timeout 120s;',
        '    }',
        '}',
    ])

    # ── 8. DEPLOYMENT CHECKLIST ──────────────────────────────────────────────
    doc.add_paragraph()
    h1(doc, '8.  Deployment Checklist for IT')

    p = doc.add_paragraph()
    r = p.add_run('Server & Software')
    r.bold = True; r.font.color.rgb = D_DARK_BLUE

    for item in [
        'Provision server with the specifications in Section 2',
        'Install Python 3.11+, pip, and ODBC Driver 17 for SQL Server',
        'Install Node.js 18+ (for the one-time frontend build)',
        'Install Nginx',
        'Install NSSM (Windows only) to run FastAPI as a Windows Service',
    ]:
        checklist_item(doc, item)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Network & DNS')
    r.bold = True; r.font.color.rgb = D_DARK_BLUE

    for item in [
        'Create subdomain paneliq.borderlessaccess.com and point to server IP',
        'Issue and install SSL certificate for the subdomain',
        'Configure HTTP → HTTPS redirect (port 80 → 443)',
        'Open outbound firewall rule: server → 10.20.30.12:1433 (SQL Server)',
        'Open outbound firewall rule: server → api.anthropic.com:443 (Claude AI API)',
    ]:
        checklist_item(doc, item)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Databases')
    r.bold = True; r.font.color.rgb = D_DARK_BLUE

    for item in [
        'Confirm KpiReports read-only access (KPIUser) is reachable from the new server',
        'Create PanelIQ_Users database with a dedicated read/write SQL login',
        'Share new DB server address and credentials with BA Analytics team',
    ]:
        checklist_item(doc, item)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Handover')
    r.bold = True; r.font.color.rgb = D_DARK_BLUE

    for item in [
        'Provide server access (RDP or SSH) to BA Analytics team for application deployment',
        'Confirm environment variables have been set on the server',
        'Share SSL certificate paths with BA Analytics team for Nginx configuration',
    ]:
        checklist_item(doc, item)

    # ── 9. QUESTIONS FOR IT ──────────────────────────────────────────────────
    doc.add_paragraph()
    h1(doc, '9.  Questions for IT')
    questions = [
        ('OS Preference',
         'Is Windows Server or Linux (Ubuntu) preferred for this deployment?'),
        ('Database Server',
         'Should the PanelIQ_Users database be created on the same SQL Server instance as KpiReports (10.20.30.12), or a separate server?'),
        ('Access Scope',
         'Should paneliq.borderlessaccess.com be accessible internally only, or also externally (VPN-only vs public)?'),
        ('SSL Certificates',
         'Who manages SSL certificates at BA — IT team or a specific infrastructure/security team?'),
        ('Firewall / Proxy',
         'Is there a web proxy or firewall that would block outbound HTTPS to api.anthropic.com? This must be confirmed before deployment.'),
        ('Server Provisioning Timeline',
         'What is the expected timeline for provisioning the server and creating the subdomain?'),
    ]
    for i, (q, detail) in enumerate(questions, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'{i}.  {q}')
        r.bold = True; r.font.color.rgb = D_DARK_BLUE; r.font.size = Pt(11)
        p2 = doc.add_paragraph(detail)
        p2.paragraph_format.left_indent = Inches(0.3)
        for run in p2.runs:
            run.font.size = Pt(10.5)
        doc.add_paragraph()

    doc.save(path)
    print(f'  Document saved: {path}')


if __name__ == '__main__':
    out = os.path.join(BASE, 'PanelIQ_IT_Requirements_June2026.docx')
    create_word(out)
    print('Done.')
