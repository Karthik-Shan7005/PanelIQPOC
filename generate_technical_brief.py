"""
Generates PanelIQ_Technical_Brief.docx
Run: python generate_technical_brief.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colours ──────────────────────────────────────────────────────────────────
BA_BLUE   = RGBColor(0x00, 0x5B, 0xAC)   # Borderless Access brand blue
BA_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
CYAN      = RGBColor(0x00, 0x94, 0xD4)
PURPLE    = RGBColor(0x6C, 0x4A, 0xD7)
LIGHT_BG  = RGBColor(0xF0, 0xF6, 0xFF)
MID_GREY  = RGBColor(0x60, 0x60, 0x70)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREEN     = RGBColor(0x00, 0x8A, 0x5E)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=BA_BLUE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)
    return p

def body(doc, text, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def inline_code(doc, label, code_text):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}  ")
    r1.font.size = Pt(10)
    r2 = p.add_run(code_text)
    r2.font.name = 'Consolas'
    r2.font.size = Pt(9)
    r2.font.color.rgb = PURPLE
    p.paragraph_format.space_after = Pt(3)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '005BAC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def two_col_table(doc, rows, header=None, col_widths=(2.5, 4.0)):
    table = doc.add_table(rows=len(rows) + (1 if header else 0), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if header:
        for i, h in enumerate(header):
            cell = table.cell(0, i)
            cell.width = Inches(col_widths[i])
            set_cell_bg(cell, '005BAC')
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
    for ri, (a, b) in enumerate(rows):
        row_idx = ri + (1 if header else 0)
        c0 = table.cell(row_idx, 0)
        c1 = table.cell(row_idx, 1)
        c0.width = Inches(col_widths[0])
        c1.width = Inches(col_widths[1])
        if ri % 2 == 0:
            set_cell_bg(c0, 'F0F6FF')
            set_cell_bg(c1, 'F0F6FF')
        r0 = c0.paragraphs[0].add_run(a)
        r0.font.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = BA_DARK
        r1 = c1.paragraphs[0].add_run(b)
        r1.font.size = Pt(9)
    doc.add_paragraph()

def callout_box(doc, title, text, bg='E8F4FD', border='005BAC'):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    p1 = cell.add_paragraph()
    r1 = p1.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = BA_BLUE
    p2 = cell.add_paragraph(text)
    p2.runs[0].font.size = Pt(9)
    doc.add_paragraph()

# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover block ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PanelIQ — Technical Brief')
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = BA_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Architecture · LLM & AI Concepts · Technology Stack')
r2.font.size = Pt(12)
r2.font.color.rgb = MID_GREY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(f'Borderless Access  ·  Confidential  ·  {datetime.date.today().strftime("%B %Y")}')
r3.font.size = Pt(10)
r3.font.color.rgb = MID_GREY

divider(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '1. Executive Summary')
body(doc,
    'PanelIQ is an AI-powered natural language interface that lets business users query the '
    'KpiReports SQL Server database in plain English — with no SQL knowledge required. '
    'It combines a FastAPI Python backend, a React web frontend, and Anthropic\'s Claude large '
    'language model to translate questions into safe, validated T-SQL queries, execute them '
    'against live panel data, and return results with charts and analyst-grade summaries.'
)
body(doc,
    'Two delivery formats are maintained in parallel:'
)
bullet(doc, 'Web Application — FastAPI + React, hosted on an internal server; the long-term production target.')
bullet(doc, 'Desktop Demo App — Electron wrapper packaging the same stack as a standalone Windows .exe; used for stakeholder demos without server dependency.')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '2. System Architecture')

heading(doc, '2.1 High-Level Architecture', level=2)
body(doc,
    'PanelIQ follows a three-tier architecture with an AI layer inserted between the '
    'presentation and data tiers:'
)

two_col_table(doc, [
    ('Presentation Tier',  'React 19 (web) / Electron shell (desktop) — user chat interface, charts, tables'),
    ('AI / API Tier',      'FastAPI (Python) — receives questions, calls Claude API, validates and executes SQL, formats response'),
    ('Data Tier',          'SQL Server — KpiReports database at 10.20.30.12; read-only via three approved views'),
], header=['Layer', 'Technology & Role'])

heading(doc, '2.2 Request Flow (per question)', level=2)
steps = [
    ('1  User types question', 'React frontend sends POST /ask to FastAPI'),
    ('2  SQL Generation',      'FastAPI calls Claude API (Call 1) with schema context + business rules as system prompt'),
    ('3  SQL Validation',      'Validator strips markdown fences, checks SELECT-only, scans for forbidden keywords, verifies approved views and global exclusion filter'),
    ('4  Query Execution',     'pyodbc executes validated T-SQL against KpiReports; results returned as rows + columns'),
    ('5  Result Formatting',   'Pandas cleans data; dates serialised; numeric columns rounded'),
    ('6  AI Summary',          'FastAPI calls Claude API (Call 2) to generate an analyst narrative and follow-up question suggestions'),
    ('7  Chart Recommendation','Logic inspects column types and row count to select bar / line / pie / metric card'),
    ('8  Response to UI',      'JSON payload returned; React renders table, chart, and summary in the chat thread'),
]
two_col_table(doc, steps, header=['Step', 'What Happens'], col_widths=(2.0, 4.5))

heading(doc, '2.3 Prompt Suggestions (Debounced — Call 3)', level=2)
body(doc,
    'As the user types (after 3+ characters, with a 600 ms debounce), the frontend calls '
    'POST /suggest. FastAPI sends the partial input to Claude (Call 3) which returns '
    '3 contextually relevant question completions. This is a lightweight call — low token '
    'count, fast response — designed to guide non-technical users toward well-formed questions.'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '3. Technology Stack')

heading(doc, '3.1 Backend', level=2)
two_col_table(doc, [
    ('FastAPI',         'High-performance Python web framework; async-ready; auto-generates OpenAPI docs at /docs'),
    ('Uvicorn',         'ASGI server that runs FastAPI; supports hot-reload in development'),
    ('Pydantic v2',     'Request/response validation and serialisation; enforces strict input types'),
    ('Anthropic SDK',   'Official Python client for the Claude API; manages authentication, retries, and streaming'),
    ('pyodbc',          'ODBC bridge for SQL Server connectivity; executes parameterised T-SQL'),
    ('Pandas',          'DataFrame-based result cleaning; handles nulls, date formatting, and type coercion'),
    ('python-dotenv',   'Loads .env credentials (API key, DB connection string) at runtime; never committed to git'),
], header=['Library', 'Role'])

heading(doc, '3.2 Frontend (Web)', level=2)
two_col_table(doc, [
    ('React 19',        'Component-based UI framework; manages chat state, message history, and loading states'),
    ('Vite 8',          'Next-generation build tool; sub-second HMR in development; optimised production bundles'),
    ('Recharts',        'Declarative chart library built on D3; renders bar, line, and pie charts from query results'),
    ('Lucide React',    'Consistent icon set used throughout the UI'),
    ('Fetch API',       'Native browser API used for REST calls to FastAPI; no additional HTTP library required'),
], header=['Library', 'Role'])

heading(doc, '3.3 Desktop App (Electron)', level=2)
two_col_table(doc, [
    ('Electron',        'Chromium + Node.js wrapper that packages web apps as native Windows/Mac/Linux desktop apps'),
    ('electron-builder','Packages the app into a distributable .exe installer (Windows) or .dmg (Mac)'),
    ('IPC Bridge',      'Electron\'s inter-process communication layer connecting the UI (renderer) to backend processes (main)'),
    ('Bundled Python',  'PyInstaller compiles the FastAPI backend into a standalone .exe; Electron spawns it as a child process'),
], header=['Component', 'Role'])

heading(doc, '3.4 Database', level=2)
two_col_table(doc, [
    ('SQL Server',          'Microsoft relational database; KpiReports instance at 10.20.30.12:1433'),
    ('ODBC Driver 17',      'Microsoft\'s SQL Server ODBC driver; required on every client machine'),
    ('3 Approved Views',    'KPISurveyData · KPIReportProjectData · KPIsurveydataExternal — only these are queryable'),
    ('Read-Only Access',    'DB user KPIUser has SELECT-only permissions; INSERT/UPDATE/DELETE blocked at DB level'),
    ('T-SQL',               'Microsoft\'s SQL dialect; PanelIQ generates and executes T-SQL exclusively'),
], header=['Component', 'Detail'])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. LLM & AI CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '4. LLM & AI Concepts Used')

heading(doc, '4.1 Large Language Model (LLM)', level=2)
body(doc,
    'PanelIQ uses Claude Sonnet 4.6 by Anthropic — a transformer-based large language model '
    'with hundreds of billions of parameters. Unlike traditional rule-based NLP, LLMs understand '
    'context, intent, and domain terminology through pre-training on vast text corpora. '
    'Claude is particularly strong at code and structured data generation tasks, making it '
    'well-suited for Text-to-SQL.'
)

heading(doc, '4.2 Natural Language to SQL (NL2SQL / Text-to-SQL)', level=2)
body(doc,
    'The core AI task in PanelIQ is NL2SQL — translating a plain English question into '
    'a syntactically correct and semantically accurate SQL query. This is one of the most '
    'studied applied NLP tasks. PanelIQ implements it using:'
)
bullet(doc, 'Schema-augmented prompting — the full database schema (column names, types, view relationships) is injected into every Claude call as context')
bullet(doc, 'Business rule injection — domain-specific rules (global exclusion filter, metric definitions, internal vs external panel logic) are embedded in the system prompt')
bullet(doc, 'Zero-shot prompting — no labelled question→SQL examples are provided; Claude infers the correct SQL purely from the schema description and rules')
body(doc,
    'This approach avoids the need for a fine-tuned model or a training dataset — the schema '
    'itself serves as the knowledge base.'
)

heading(doc, '4.3 Prompt Engineering', level=2)
body(doc,
    'Prompt engineering is the practice of carefully designing the text input to an LLM to '
    'steer its output. PanelIQ uses a structured system prompt (~230 lines) that contains:'
)
two_col_table(doc, [
    ('Role definition',      '"You are a SQL Server expert for a market research panel analytics platform"'),
    ('Schema context',       'Full column-level description of all 3 views with data types and nullability'),
    ('Business rules',       '8 numbered rules: global exclusion, date handling, metric formulas, status groups, panel source logic'),
    ('Date translation',     'Mapping of natural language time expressions to T-SQL date functions'),
    ('Standard JOIN patterns','Pre-defined JOIN templates to avoid hallucinated table relationships'),
    ('Output rules',         '"Return ONLY the SQL — no explanation, no markdown, no backticks"'),
], header=['Section', 'Content'], col_widths=(2.2, 4.3))

heading(doc, '4.4 System Prompt vs User Prompt', level=2)
body(doc,
    'Claude (and most modern LLMs) accept two input channels per API call:'
)
bullet(doc, 'System prompt — persistent instructions that define the model\'s role, constraints, and knowledge. In PanelIQ this is the ~230-line schema context. It is sent with every request.')
bullet(doc, 'User prompt — the actual question typed by the user (e.g. "Show completes by market for March 2026"). This changes every request.')
body(doc,
    'Separating these allows the expensive schema context to be cached (Anthropic\'s prompt '
    'caching feature), reducing API cost and latency on repeated calls.'
)

heading(doc, '4.5 Context Window & Token Management', level=2)
body(doc,
    'LLMs process text as tokens (roughly 0.75 words per token). Claude Sonnet 4.6 supports '
    'a 200,000-token context window. PanelIQ manages tokens as follows:'
)
two_col_table(doc, [
    ('SQL Generation call', 'System: ~1,800 tokens (schema). User: ~20–50 tokens. Output: ~300–800 tokens.'),
    ('Summary call',        'System: short role prompt. User: question + SQL + result rows. Output: ~200 tokens.'),
    ('Suggestion call',     'System: short role prompt. User: partial input. Output: ~60 tokens (3 suggestions).'),
    ('max_tokens setting',  '1,000 tokens capped on SQL generation to prevent runaway output'),
], header=['Call', 'Token Profile'], col_widths=(2.2, 4.3))

heading(doc, '4.6 Hallucination & Guardrails', level=2)
body(doc,
    'LLM hallucination — generating plausible but incorrect output — is a key risk in NL2SQL. '
    'PanelIQ mitigates this with a multi-layer defence:'
)
two_col_table(doc, [
    ('Structural validation',   'SELECT-only check; strips markdown fences; rejects empty queries'),
    ('Keyword blocklist',       'Scans for INSERT, UPDATE, DELETE, DROP, EXEC, ALTER, XP_ etc.'),
    ('View allowlist',          'Query must reference at least one of the 3 approved views'),
    ('Business rule check',     'Query must contain the global ClientName exclusion filter'),
    ('DB-level enforcement',    'KPIUser account has SELECT-only permission — even if validation fails, write operations are blocked at the database'),
], header=['Layer', 'What It Catches'], col_widths=(2.5, 4.0))

heading(doc, '4.7 Retrieval-Augmented Generation (RAG) — Concept', level=2)
body(doc,
    'PanelIQ uses a simplified form of RAG: instead of retrieving documents at runtime, '
    'the schema and business rules are statically embedded in the system prompt — a pattern '
    'known as static RAG or schema grounding. A future enhancement could implement dynamic '
    'RAG, retrieving only the relevant schema sections based on the user\'s question, reducing '
    'token cost as the database grows.'
)

heading(doc, '4.8 Chain of Thought Pipeline', level=2)
body(doc,
    'PanelIQ implements an explicit multi-step reasoning chain rather than asking Claude to '
    'produce everything in one call. This improves quality and allows each step to be '
    'validated independently:'
)
bullet(doc, 'Step 1 — Generate SQL (Claude, Call 1)')
bullet(doc, 'Step 2 — Validate SQL (deterministic Python rules, no AI)')
bullet(doc, 'Step 3 — Execute SQL (SQL Server, no AI)')
bullet(doc, 'Step 4 — Summarise results (Claude, Call 2)')
body(doc,
    'This separation means a bad SQL query is caught before execution — Claude is not asked '
    'to both generate and validate its own output, which would be unreliable.'
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. SECURITY ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '5. Security Architecture')

two_col_table(doc, [
    ('API Key management',      '.env file gitignored; key loaded at runtime via python-dotenv; never hardcoded'),
    ('SQL injection prevention','All SQL is Claude-generated and validated; no user input is concatenated into SQL strings'),
    ('Read-only database user', 'KPIUser has SELECT-only grants on KpiReports; no write operations possible'),
    ('SELECT-only enforcement', 'Dual enforcement: application validator + database permissions'),
    ('Global exclusion filter', '\'Engagement Surveys - Research\' excluded from every query at the AI prompt level'),
    ('CORS policy',             'Backend accepts requests only from localhost:5173 and localhost:3000 (dev); production will restrict to the internal domain'),
    ('No PII in logs',          'Query logs print only row counts and chart type — no respondent data'),
], header=['Control', 'Implementation'], col_widths=(2.5, 4.0))
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. WEB APPLICATION vs DESKTOP APP — COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '6. Web Application vs Desktop Demo App')

two_col_table(doc, [
    ('Delivery',            'Browser URL (internal network)',               'Windows .exe installer'),
    ('Frontend',            'React 19 served by Vite / Nginx',             'Same React app inside Electron Chromium'),
    ('Backend',             'FastAPI on internal server (uvicorn)',         'FastAPI compiled by PyInstaller, spawned by Electron'),
    ('Database access',     'Server connects to SQL Server',               'User\'s machine connects to SQL Server (VPN required)'),
    ('Updates',             'Deploy to server — all users get it instantly','Redistribute new .exe installer'),
    ('Tech team needed',    'Yes — for server provisioning and hosting',   'No — self-contained installer'),
    ('Multi-user support',  'Yes — concurrent users, shared history',      'Single user per install'),
    ('Best for',            'Production deployment (Phase 1+)',             'COO demo, offline presentations'),
], header=['Aspect', 'Web Application', 'Desktop Demo App'], col_widths=(2.0, 2.6, 2.6))[0] if False else None

# manual 3-col table
table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
headers = ['Aspect', 'Web Application', 'Desktop Demo App']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    set_cell_bg(cell, '005BAC')
    run = cell.paragraphs[0].add_run(h)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9)

rows = [
    ('Delivery',           'Browser URL on internal network',            'Windows .exe installer'),
    ('Frontend',           'React 19 served by Vite / Nginx',            'Same React app inside Electron Chromium'),
    ('Backend',            'FastAPI on a server (uvicorn)',               'FastAPI compiled via PyInstaller, spawned by Electron'),
    ('Database access',    'Server → SQL Server (always connected)',      'User machine → SQL Server (VPN required)'),
    ('Updates',            'Deploy once; all users get it instantly',     'Redistribute new .exe installer'),
    ('Tech team needed',   'Yes — server provisioning & hosting',        'No — fully self-contained'),
    ('Multi-user',         'Yes — concurrent sessions supported',        'Single user per install'),
    ('Best for',           'Production (Phase 1+)',                       'COO demo, offline presentations'),
]
for ri, (a, b, c) in enumerate(rows):
    row = table.rows[ri + 1]
    bg = 'F0F6FF' if ri % 2 == 0 else 'FFFFFF'
    for ci, cell in enumerate(row.cells):
        set_cell_bg(cell, bg)
    ra = row.cells[0].paragraphs[0].add_run(a)
    ra.font.bold = True; ra.font.size = Pt(9)
    row.cells[1].paragraphs[0].add_run(b).font.size = Pt(9)
    row.cells[2].paragraphs[0].add_run(c).font.size = Pt(9)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. API DESIGN
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '7. API Design')

two_col_table(doc, [
    ('GET  /',          'Health check — confirms API is running'),
    ('GET  /health',    'Detailed health check — returns DB name, AI model, module name'),
    ('POST /ask',       'Main endpoint — accepts {"question": "..."}, returns SQL + data + summary + chart type'),
    ('POST /suggest',   'Autocomplete endpoint — accepts {"partial": "..."}, returns 3 question suggestions'),
], header=['Endpoint', 'Purpose'])

callout_box(doc,
    'REST API — Key Design Principles',
    'Stateless (each request is self-contained) · JSON in/out · Pydantic-validated request bodies · '
    'CORS-restricted to known origins · Auto-documented at /docs via FastAPI\'s built-in OpenAPI support',
    bg='E8F4FD'
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. FUTURE ML / AI ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '8. Future AI / ML Enhancements')

two_col_table(doc, [
    ('Prompt caching',          'Anthropic\'s cache_control feature pins the schema system prompt in cache for 5 minutes; reduces cost by ~90% on repeated calls'),
    ('Dynamic RAG',             'As the schema grows (Phase 2 vendor data, Phase 3 finance), retrieve only relevant schema sections per question rather than sending everything'),
    ('Query history & feedback','Store past questions + SQL in a second database; use accepted queries as few-shot examples to improve accuracy over time'),
    ('Fine-tuning (future)',     'With enough validated question→SQL pairs, fine-tune a smaller model on BA-specific data to reduce Claude API dependency'),
    ('Anomaly detection',       'Apply statistical ML (isolation forest, z-score) to flag unusual panel metrics automatically — proactive alerts rather than reactive queries'),
    ('Semantic search',         'Embed past questions with vector embeddings; surface similar historical queries to help users discover relevant analyses'),
    ('Row-level access control','Inject user identity into the system prompt; restrict SQL to user\'s business unit data at the AI layer, not just DB permissions'),
], header=['Enhancement', 'Description'], col_widths=(2.5, 4.0))
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 9. KEY TERMINOLOGY REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '9. Key Terminology Reference')

two_col_table(doc, [
    ('LLM',             'Large Language Model — AI model trained on vast text to understand and generate human language (e.g. Claude, GPT-4)'),
    ('NL2SQL / Text-to-SQL', 'Converting a plain English question into a SQL database query automatically'),
    ('Prompt Engineering','Crafting precise instructions for an LLM to control its output — the core skill in building AI applications without model training'),
    ('System Prompt',   'Persistent instructions given to an LLM before the user\'s message; defines role, rules, and knowledge'),
    ('Zero-shot',       'Asking the LLM to perform a task with no examples provided — relies entirely on the model\'s pre-trained knowledge and the system prompt'),
    ('Few-shot',        'Providing 2–5 worked examples in the prompt to guide the model; improves accuracy on niche tasks'),
    ('RAG',             'Retrieval-Augmented Generation — injecting relevant external knowledge into the prompt at query time so the model can answer questions beyond its training data'),
    ('Hallucination',   'When an LLM generates confident but factually incorrect output — mitigated in PanelIQ via SQL validation and read-only DB access'),
    ('Token',           'The unit of text an LLM processes (~0.75 English words). API cost and context limits are measured in tokens'),
    ('Context Window',  'Maximum tokens an LLM can process in one call. Claude Sonnet 4.6 supports 200,000 tokens'),
    ('Transformer',     'The neural network architecture underlying modern LLMs; enables parallel processing of entire sequences and long-range context understanding'),
    ('API',             'Application Programming Interface — the communication contract between the frontend, backend, and Claude; PanelIQ uses REST APIs'),
    ('FastAPI',         'Python web framework for building APIs; chosen for speed, automatic validation, and built-in API documentation'),
    ('Electron',        'Framework to build cross-platform desktop apps using web technologies (HTML/CSS/JS); wraps PanelIQ for offline demo delivery'),
    ('Pydantic',        'Python library for data validation; ensures request payloads match expected types before processing'),
    ('ODBC',            'Open Database Connectivity — standard interface for connecting applications to databases; pyodbc provides this bridge to SQL Server'),
    ('T-SQL',           'Transact-SQL — Microsoft\'s SQL dialect used by SQL Server; includes stored procedures, window functions, and date functions not in standard SQL'),
], header=['Term', 'Plain English Definition'], col_widths=(2.0, 4.5))

# ── Footer note ───────────────────────────────────────────────────────────────
divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    f'PanelIQ Technical Brief  ·  Borderless Access  ·  Confidential  ·  '
    f'{datetime.date.today().strftime("%d %B %Y")}  ·  '
    f'Built with FastAPI · React · Claude Sonnet 4.6 · SQL Server'
)
r.font.size = Pt(8)
r.font.color.rgb = MID_GREY

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = 'PanelIQ_Technical_Brief_June2026.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
